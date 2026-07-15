# -*- coding: utf-8 -*-
"""
应用安全检测报告 · 标准模板生成器
=================================
按客户《应用安全检测报告》标准模板(13 类 / 55 项)生成 DOCX。

用法:
    python3 gen_report.py --items items_55.json --findings findings.json \
            --meta meta.json --evidence ./evidence --out 报告.docx

真实性原则(必须遵守):
    * 每项 测试结果 必须对应真实证据(HTTP 请求/响应、页面、API、前端 JS、截图或本地证据文件);
    * 未做的测试一律标 "未测试" 并在 截图说明 里写明原因;
    * 设计能力/架构隐患标 "设计性风险，非本次实际利用结果",不得写成已利用;
    * 报告中不得出现口令、完整 Cookie、令牌/PAT、AK/SK、32 位散列等敏感原文。

输入文件:
    items_<n>.json  标准检测项清单(环境无关): 列表, 每项含
        测试编号/风险名称/风险等级/功能要求/风险影响/测试依据
    findings.json   本次实测结论(按编号): { "PORT-001": {
        "测试结果": "不符合|符合|未测试|不涉及",
        "截图说明": "...", "修复建议": "...",
        "截图": "fig_xxx.png"(可空), "图注": "图1: ...", "备注": "" }, ... }
        —— 必须覆盖 items 中的每一个测试编号。
    meta.json       报告元信息: {
        "项目名称","副标题"(如"初测"),"委托单位","测试人员","测试日期",
        "编制单位","应用网址","系统级别","IP地址",
        "参考标准":[...], "检测工具":[["名称","类型","用途"],...],
        "变更记录":[["时间","版本","说明","修改人"],...] }
"""
import json, os, argparse
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CATEGORIES = [
    ('PORT', '端口安全'), ('BASIC', '基线安全'), ('FRON', '前端安全'),
    ('TRAN', '传输安全'), ('AUTH', '身份鉴别'), ('PER', '权限管理'),
    ('SESS', '会话安全'), ('Key', '密钥安全'), ('INFO', '异常处理'),
    ('BUS', '业务安全'), ('AUD', '日志审计'), ('DATA', '数据安全'),
    ('VULN', '常见漏洞'),
]
RESULT_COLOR = {'不符合': RGBColor(0xC0, 0x00, 0x00), '符合': RGBColor(0x00, 0x80, 0x00),
                '未测试': RGBColor(0xBF, 0x8F, 0x00), '不涉及': RGBColor(0x59, 0x59, 0x59)}
LEVEL_COLOR = {'严重': RGBColor(0xC0, 0x00, 0x00), '高危': RGBColor(0xE3, 0x6C, 0x09),
               '高': RGBColor(0xE3, 0x6C, 0x09), '中危': RGBColor(0xBF, 0x8F, 0x00),
               '中': RGBColor(0xBF, 0x8F, 0x00), '低危': RGBColor(0x00, 0x70, 0xC0),
               '低': RGBColor(0x00, 0x70, 0xC0)}
# 详表/汇总按模板以单字展示风险等级
LEVEL_SHORT = {'严重': '严重', '高危': '高', '高': '高', '中危': '中', '中': '中',
               '低危': '低', '低': '低'}
# 统计口径(把单字/长词都折算到长词)
LEVEL_FULL = {'严重': '严重', '高': '高危', '高危': '高危', '中': '中危', '中危': '中危',
              '低': '低危', '低危': '低危'}


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def build(items_meta, findings, meta, evidence_dir, out_path):
    order = {c[0]: i for i, c in enumerate(CATEGORIES)}

    def prefix(idv):
        return idv.split('-')[0]

    # 组装条目并校验覆盖
    meta_by_id = {m['测试编号']: m for m in items_meta}
    missing = [i for i in meta_by_id if i not in findings]
    if missing:
        raise SystemExit('findings 缺少以下测试编号: ' + '、'.join(sorted(missing)))
    items = []
    for idv, m in meta_by_id.items():
        f = findings[idv]
        lvl = f.get('风险等级', '').strip()
        items.append((idv, m, f.get('测试结果', ''), f.get('截图说明', ''),
                      f.get('修复建议', ''), f.get('截图') or None,
                      f.get('图注', ''), f.get('备注', ''), lvl))
    items.sort(key=lambda x: (order[prefix(x[0])], x[0]))

    result_counts = Counter(x[2] for x in items)
    fail_levels = Counter(LEVEL_FULL.get(x[8], x[8])
                          for x in items if x[2] == '不符合')
    n = len(items)
    n_fail = result_counts['不符合']; n_ok = result_counts['符合']
    n_ut = result_counts['未测试']; n_na = result_counts['不涉及']
    sev = fail_levels.get('严重', 0); high = fail_levels.get('高危', 0)
    mid = fail_levels.get('中危', 0); low = fail_levels.get('低危', 0)

    # 章节号 3.<类序>.<类内序>
    chapter = {}; cat_seen = []; cat_idx = {}
    for tup in items:
        idv = tup[0]
        cat = prefix(idv)
        if cat not in cat_seen:
            cat_seen.append(cat); cat_idx[cat] = 0
        cat_idx[cat] += 1
        chapter[idv] = f'3.{cat_seen.index(cat) + 1}.{cat_idx[cat]}'

    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)

    def center(txt, sz, bold=True, color=None):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(sz); r.bold = bold
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            r.font.color.rgb = color
        return p

    def h1(t): doc.add_heading(t, level=1)
    def h2(t): doc.add_heading(t, level=2)
    def h3(t): doc.add_heading(t, level=3)

    def para(t, italic=False):
        p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
        return p

    def table_from(rows, cols, style='Light Grid Accent 1'):
        t = doc.add_table(rows=0, cols=cols); t.style = style
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = str(v)
                if i == 0:
                    for x in cells[j].paragraphs[0].runs:
                        x.bold = True
        return t

    # ---------- 封面 ----------
    for _ in range(4):
        doc.add_paragraph()
    center(meta['项目名称'], 26)
    center('应用安全检测报告', 26)
    for _ in range(2):
        doc.add_paragraph()
    center(meta.get('副标题', '初测'), 16, color=RGBColor(0xC0, 0x00, 0x00))
    for _ in range(6):
        doc.add_paragraph()
    for line in ['委托单位：' + meta.get('委托单位', ''),
                 '测试人员：' + meta.get('测试人员', ''),
                 '测试日期：' + meta.get('测试日期', '')]:
        center(line, 13, bold=False)
    for _ in range(6):
        doc.add_paragraph()
    center(meta.get('编制单位', ''), 12, bold=False)
    doc.add_page_break()

    # ==================== 1 检测概述 ====================
    h1('检测概述')
    h2('背景介绍')
    para('随着国内外网络安全形势日益严峻，攻防不对等拉锯，网络攻击入侵门槛不断拉低，同时网络安全法的落地实施，'
         '国家监管层面不断加大信息系统安全监管及处罚力度。一旦信息系统应用软件存在安全漏洞、配置不规范或业务'
         '安全风险，易遭受到网络安全攻击，给应用系统的稳定运行、数据安全带来了极大的风险隐患，信息系统的安全'
         '已经成为了目前迫切亟待解决的问题。')
    para('为进一步加强网络安全管理与数据安全保护，切实防范信息安全风险，依据国家（行业）标准和监管合规要求，'
         '对目标系统 ' + meta.get('应用网址', '') + ' 开展应用安全检测，全面识别应用系统面临的安全风险隐患，'
         '包含业务安全风险，及时采取相应的技术措施进行安全防范和修复加固，降低攻击入侵，提升业务系统安全健壮性，'
         '保障业务系统安全运行。')

    h2('参考标准')
    para('此检测报告引用以下标准与规范（版本于报告编制时经官方来源核对，均为现行有效版本，并结合被测对象'
         '所属行业选取相关项）：')
    for s in meta.get('参考标准', []):
        doc.add_paragraph(s, style='List Bullet')
    para('注：上述标准每次执行检测前应重新核对是否有更新版本。', italic=True)

    h2('适用说明')
    para('针对软件应用（含移动服务端、C/S 服务端等）开展应用安全检测，本次被测对象为 '
         + meta['项目名称'] + '，不涉及主机、配置等其他检测。')

    h2('工作原则')
    para('开展安全检测应遵循的工作原则是：')
    for k, v in [
        ('（1）规范性原则', '安全检测过程有统一的流程，安全检测中使用的方法及使用的文档，需要具有很好的规范性，'
                       '便于安全检测工作的质量保证，并保证结果的一致性。'),
        ('（2）可控性原则', '安全检测所使用的工具、方法和过程要在双方认可的范围之内，安全检测的进度要符合进度表的'
                       '安排，保证双方对安全检测的可控性。'),
        ('（3）整体性原则', '安全检测的范围和内容应当系统性、全面，覆盖信息系统安全所涉及的各个层面，并考虑各个'
                       '层面的相互关系。'),
        ('（4）最小影响原则', '安全检测应尽可能较小的影响网络和系统的正常运行，不能对系统的业务产生显著影响，'
                        '避免造成被测评系统的性能明显下降、网络拥塞、服务中断等；不删除或修改数据、不触发'
                        '破坏性操作、不进行高强度爆破、不执行破坏性注入类测试。'),
        ('（5）保密性原则', '对在安全检测中所接触到的用户单位的所有敏感信息，检测人员应遵循相关的保密承诺，不利用'
                       '相关信息进行任何侵害被检测单位安全利益的行为。本报告不含口令、完整 Cookie、令牌等敏感原文。'),
    ]:
        doc.add_paragraph(k)
        para(v)

    h2('检测内容')
    para('该检测报告用于应用系统安全检测，可覆盖已上线前安全检测、上线后安全检测、安全评估、互联网安全检测等'
         '场景，包含端口安全、基线安全、前端安全、传输安全、身份鉴别、权限管理、会话安全、密钥安全、异常处理、'
         '业务安全、日志审计、数据安全、常见漏洞等 13 个大类，共计 ' + str(n) + ' 项安全检测内容。')

    h2('检测工具')
    para('安全检测工具清单如下（当前列举为本次实际使用的主要工具，结合不同应用类型和漏洞组件验证，会使用到其他'
         '检测工具）：')
    table_from([('工具名称', '工具类型', '工具用途')] + [tuple(r) for r in meta.get('检测工具', [])], 3)

    h2('检测声明')
    for s in ['安全检测在应用系统上线前开展，如测试环境和生产环境一致，宜使用测试环境开展，如测试和生产差异较大，'
              '两套环境应同时开展检测，检测前应确保应用所有组件部署完成，并提供生产和测试环境所有资产信息。',
              '本报告为本次应用安全检测服务的输出结果，供业务方进行漏洞修补和安全加固参考，并进行安全风险规避，'
              '安全控制和管理。',
              '本报告的检测结论仅对被测对象当时的状态有效。当检测完成后，由于被测对象的系统变更，本报告不再适用。',
              '在任何情况下，若需引用本报告中的结果或结论都应保持其原有的意义，不得对相关内容擅自进行增加、修改和'
              '伪造或掩盖事实。']:
        para(s)

    h2('变更记录')
    table_from([('时间', '版本', '说明', '修改人')]
               + [tuple(r) for r in meta.get('变更记录', [['', 'V1.0', '创建应用安全检测报告', '']])], 4)

    # ==================== 2 检测结果 ====================
    h1('检测结果')
    h2('检测信息')
    table_from([('项目名称', '系统级别', '应用网址', 'IP地址'),
                (meta['项目名称'], meta.get('系统级别', ''),
                 meta.get('应用网址', ''), meta.get('IP地址', ''))], 4)

    h2('风险信息')
    para(f'本次安全检测共完成 {n} 项检测内容，共计发现 {n_fail} 个安全风险，其中严重风险 {sev} 个、'
         f'高危风险 {high} 个、中危风险 {mid} 个、低危风险 {low} 个（判定为“不符合”）。'
         f'另有“符合” {n_ok} 项、“未测试” {n_ut} 项、“不涉及” {n_na} 项。')
    para(f'“未测试” {n_ut} 项系因测试条件受限（如缺少对照账号、生产环境非破坏性约束、经隧道无法验证等）'
         f'无法完成验证，据实标记并说明原因；“不涉及” {n_na} 项系统不涉及对应功能场景。')

    h2('风险汇总')
    t = doc.add_table(rows=0, cols=5); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.add_row().cells
    for j, v in enumerate(['序号', '风险章节', '风险名称', '风险等级', '复测结果']):
        hdr[j].text = v
        for x in hdr[j].paragraphs[0].runs:
            x.bold = True
    seq = 0
    for idv, m, res, ev, fix, fig, cap, remark, lvl in items:
        if res != '不符合':
            continue
        seq += 1
        cells = t.add_row().cells
        cells[0].text = str(seq); cells[1].text = chapter[idv]
        cells[2].text = m['风险名称']
        cells[3].text = LEVEL_SHORT.get(lvl, lvl)
        cells[4].text = ''
        for run in cells[3].paragraphs[0].runs:
            run.font.color.rgb = LEVEL_COLOR.get(lvl, RGBColor(0, 0, 0))
            run.bold = True
    para(f'注：上表仅列出本次判定为“不符合”的 {n_fail} 项风险；全部 {n} 项检测内容的逐项结果'
         f'（含符合/未测试/不涉及）详见“检测详情”。“复测结果”列供整改复测时填写。', italic=True)

    # ==================== 3 检测详情 ====================
    h1('检测详情')
    cur = None
    for idv, m, res, ev, fix, fig, cap, remark, lvl in items:
        cat = prefix(idv)
        if cat != cur:
            cur = cat
            h2(dict(CATEGORIES)[cat])
        h3(f"{idv} {m['风险名称']}")
        tt = doc.add_table(rows=0, cols=2); tt.style = 'Table Grid'
        tt.alignment = WD_TABLE_ALIGNMENT.CENTER
        tt.columns[0].width = Inches(1.3); tt.columns[1].width = Inches(5.2)

        def add_kv(k, v, color=None, bold=False):
            c = tt.add_row().cells
            c[0].text = k; c[1].text = v
            set_cell_bg(c[0], 'F2F2F2')
            for r_ in c[0].paragraphs[0].runs:
                r_.bold = True
            if color:
                for r_ in c[1].paragraphs[0].runs:
                    r_.font.color.rgb = color; r_.bold = bold

        add_kv('测试编号', idv)
        add_kv('风险名称', m['风险名称'])
        add_kv('风险等级', LEVEL_SHORT.get(lvl, lvl) if lvl else '—',
               LEVEL_COLOR.get(lvl), True)
        add_kv('功能要求', m['功能要求'])
        add_kv('风险影响', m['风险影响'])
        add_kv('测试依据', m['测试依据'])
        add_kv('测试结果', res, RESULT_COLOR.get(res), True)
        add_kv('截图说明', ev)
        add_kv('修复建议', fix)
        add_kv('备注', remark)
        if fig:
            fp = os.path.join(evidence_dir, fig)
            if os.path.exists(fp):
                doc.add_picture(fp, width=Inches(5.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rc = pc.add_run(cap); rc.italic = True; rc.font.size = Pt(9)
        doc.add_paragraph()

    doc.save(out_path)
    print('SAVED', out_path)
    print('items=', n, 'results=', dict(result_counts), 'fail_levels=', dict(fail_levels))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--items', required=True)
    ap.add_argument('--findings', required=True)
    ap.add_argument('--meta', required=True)
    ap.add_argument('--evidence', default='.')
    ap.add_argument('--out', required=True)
    a = ap.parse_args()
    items_meta = json.load(open(a.items, encoding='utf-8'))
    findings = json.load(open(a.findings, encoding='utf-8'))
    meta = json.load(open(a.meta, encoding='utf-8'))
    build(items_meta, findings, meta, a.evidence, a.out)


if __name__ == '__main__':
    main()
